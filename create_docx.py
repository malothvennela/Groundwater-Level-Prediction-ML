from pathlib import Path
from docx import Document

INPUT_FILE = Path("outputs/final_project_report.md")
OUTPUT_FILE = Path("outputs/Groundwater_Level_Prediction_Final_Documentation.docx")

document = Document()
document.add_heading("Groundwater Level Prediction Using Machine Learning", 0)

if not INPUT_FILE.exists():
    INPUT_FILE = Path("outputs/mentor_required_summary.md")

text = INPUT_FILE.read_text(encoding="utf-8")

for line in text.splitlines():
    line = line.strip()

    if not line:
        continue

    if line.startswith("# "):
        document.add_heading(line.replace("# ", ""), level=1)
    elif line.startswith("## "):
        document.add_heading(line.replace("## ", ""), level=2)
    elif line.startswith("### "):
        document.add_heading(line.replace("### ", ""), level=3)
    elif line.startswith("- "):
        document.add_paragraph(line.replace("- ", ""), style="List Bullet")
    elif line.startswith("|"):
        continue
    elif line.startswith("```"):
        continue
    else:
        document.add_paragraph(line)

OUTPUT_FILE.parent.mkdir(exist_ok=True)
document.save(OUTPUT_FILE)

print(f"Document created: {OUTPUT_FILE}")
